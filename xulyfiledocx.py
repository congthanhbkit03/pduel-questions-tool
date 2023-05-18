import sys
import docx
from question import Question
from texfile import xay_dung_file_tex

# docx_filename = "Test.docx"
docx_filename = sys.argv[1] #python xulyfiledoc.py bai1.docx
dau_phan_cach_cau = sys.argv[2] # co nghia la kí tự phân cách sau chữ Câu như:  Câu 1., Câu 2. hay Câu 1:, Câu 2:
myDoc = docx.Document(docx_filename)
question_arr = []
dapan_arr = []

# lay tieu de bai thi
quiz_title = myDoc.paragraphs[0].text

# lay dap an bai thi
# dapan_str = myDoc.paragraphs[len(myDoc.paragraphs) - 1].text
# # dapan_with_semicolon_arr có dạng [ '1:A', '2:D', '3:B' ....]
# dapan_with_semicolon_arr = dapan_str.split("|")
# # xu ly chi lay phan sau : => ['A', 'D', 'B' ....] => sau đó thành [0, 3, 1 ...]
# for da in dapan_with_semicolon_arr:
#     # if da[da.find(":") + 1 :].strip() == 'A':
#     #     opt_id = 0
#     # elif da[da.find(":") + 1 :].strip() == 'B':
#     #     opt_id = 1
#     # elif da[da.find(":") + 1 :].strip() == 'C':
#     #     opt_id = 2
#     # else:
#     #     opt_id = 3
#     #
#     # dapan_arr.append(opt_id)
#
#     dapan_arr.append(da[da.find(":") + 1:].strip())

#xu ly dap an voi table
table = myDoc.tables[0] #table dau tien
dapan_arr = []
for i, row in enumerate(table.rows):
    text = (cell.text for cell in row.cells)
    # print(text)
    if i % 2 == 1:
        row_data = list(text)
        # dapan_arr.append(row_data)
        dapan_arr += row_data

print(dapan_arr)
print(len(myDoc.paragraphs))
# kiem tra neu so luong dap an <> so luong cau hoi thi bao loi
if len(dapan_arr) != len(myDoc.paragraphs) - 2:
    print("Xem lai so luong dap an - So luong cau hoi")
    exit()
# print(dapan_str)

for idx, prg in enumerate(myDoc.paragraphs):
    # paragraph dau tien chua Category - Tieu de cho bai hi
    # paragraph cuoi cung chua dap an  - nen loai 2 paragraph nay ra

    if idx > 0 and idx < len(myDoc.paragraphs) - 1:
        # phan tich paragraphs va dua vao object Question

        #tim . đầu tiên => lấy phần title
        index_1 = prg.text.find(dau_phan_cach_cau)    #dinh dang Câu 1. nội dung ...
        title = prg.text[0:index_1]
        # print(title)

        # tim A.
        index_2 = prg.text.find("A.")
        content = prg.text[index_1 + 1: index_2]

        index_3 = prg.text.find("B.")
        optA = prg.text[index_2 + 2: index_3]

        index_4 = prg.text.find("C.")
        optB = prg.text[index_3 + 2: index_4]

        index_5 = prg.text.find("D.")
        optC = prg.text[index_4 + 2:index_5]

        optD = prg.text[index_5 + 2:]

        correct_i = 1

        q = Question(title, content, optA, optB, optC, optD, correct_i)
        # print(q)
        question_arr.append(q)

# tao file .text voi quiz_tile, danh sach cau hoi va danh sach dap an
xay_dung_file_tex(docx_filename, quiz_title, question_arr, dapan_arr)


